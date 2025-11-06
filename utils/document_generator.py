"""
Document Generator for Project Assistant
Creates professional Word/PDF documents for final project submission
"""

import logging
from datetime import datetime
from typing import Dict, Optional
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Document generation will be limited.")

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Handles creation of professional project documents"""
    
    def create_project_document(self, project_data: Dict) -> Optional[str]:
        """
        Create a complete ZIMSEC project document in Word format
        
        Args:
            project_data: Dictionary containing all project information
            
        Returns:
            URL or path to the generated document, or None if failed
        """
        try:
            if not DOCX_AVAILABLE:
                logger.error("python-docx not available")
                return None
            
            # Create document
            doc = Document()
            
            # Set up styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            # Title Page
            self._add_title_page(doc, project_data)
            doc.add_page_break()
            
            # Table of Contents (placeholder)
            self._add_table_of_contents(doc)
            doc.add_page_break()
            
            # Stage 1: Problem Identification
            if 'stage_1_data' in project_data:
                self._add_stage_1(doc, project_data)
                doc.add_page_break()
            
            # Stage 2: Investigation
            if 'stage_2_data' in project_data:
                self._add_stage_2(doc, project_data)
                doc.add_page_break()
            
            # Stage 3: New Ideas
            if 'stage_3_data' in project_data:
                self._add_stage_3(doc, project_data)
                doc.add_page_break()
            
            # Stage 4: Development
            if 'stage_4_data' in project_data:
                self._add_stage_4(doc, project_data)
                doc.add_page_break()
            
            # Stage 5: Presentation
            if 'stage_5_data' in project_data:
                self._add_stage_5(doc, project_data)
                doc.add_page_break()
            
            # Stage 6: Evaluation
            if 'stage_6_data' in project_data:
                self._add_stage_6(doc, project_data)
            
            # Save document
            filename = f"ZIMSEC_Project_{project_data.get('project_title', 'Document').replace(' ', '_')}.docx"
            doc_path = f"/tmp/{filename}"
            doc.save(doc_path)
            
            logger.info(f"Document created: {doc_path}")
            return doc_path
            
        except Exception as e:
            logger.error(f"Error creating document: {e}", exc_info=True)
            return None
    
    def _add_title_page(self, doc: Document, project_data: Dict):
        """Add title page to document"""
        # School name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ZIMBABWE SCHOOL EXAMINATIONS COUNCIL\n")
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Project title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_data.get('project_title', 'School-Based Project'))
        run.font.size = Pt(18)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Subject
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Subject: {project_data.get('subject', 'N/A')}")
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Student name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared by: {project_data.get('student_name', 'Student')}")
        run.font.size = Pt(12)
        
        # Date
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime("%B %Y"))
        run.font.size = Pt(11)
    
    def _add_table_of_contents(self, doc: Document):
        """Add table of contents"""
        heading = doc.add_heading('Table of Contents', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('1. Problem Identification')
        doc.add_paragraph('2. Investigation of Related Ideas')
        doc.add_paragraph('3. Generation of New Ideas')
        doc.add_paragraph('4. Development of the Best Idea')
        doc.add_paragraph('5. Presentation of Results')
        doc.add_paragraph('6. Evaluation and Recommendations')
    
    def _add_stage_1(self, doc: Document, project_data: Dict):
        """Add Stage 1 content"""
        doc.add_heading('STAGE 1: PROBLEM IDENTIFICATION', level=1)
        
        stage_data = project_data.get('stage_1_data', {})
        
        doc.add_heading('1.1 Problem Statement', level=2)
        doc.add_paragraph(stage_data.get('problem_ideas', [''])[0])
        
        doc.add_heading('1.2 Location', level=2)
        doc.add_paragraph(stage_data.get('location', 'Not specified'))
        
        doc.add_heading('1.3 Affected Population', level=2)
        doc.add_paragraph(stage_data.get('affected_people', 'Not specified'))
        
        doc.add_heading('1.4 Challenges Created', level=2)
        doc.add_paragraph(stage_data.get('challenges', 'Not specified'))
        
        doc.add_heading('1.5 Project Objectives', level=2)
        doc.add_paragraph('The primary objective of this project is to address the identified problem through systematic investigation and innovative solution development.')
    
    def _add_stage_2(self, doc: Document, project_data: Dict):
        """Add Stage 2 content"""
        doc.add_heading('STAGE 2: INVESTIGATION OF RELATED IDEAS', level=1)
        
        stage_data = project_data.get('stage_2_data', {})
        
        doc.add_heading('2.1 Literature Review', level=2)
        search_results = stage_data.get('search_results', 'No web research conducted.')
        doc.add_paragraph(search_results)
        
        doc.add_heading('2.2 Analysis of Existing Solutions', level=2)
        analysis = stage_data.get('analysis', 'Analysis pending.')
        doc.add_paragraph(analysis)
    
    def _add_stage_3(self, doc: Document, project_data: Dict):
        """Add Stage 3 content"""
        doc.add_heading('STAGE 3: GENERATION OF NEW IDEAS', level=1)
        
        stage_data = project_data.get('stage_3_data', {})
        ideas = stage_data.get('ideas', [])
        
        doc.add_heading('3.1 Brainstormed Solutions', level=2)
        for idx, idea in enumerate(ideas, 1):
            doc.add_paragraph(f"{idx}. {idea}")
    
    def _add_stage_4(self, doc: Document, project_data: Dict):
        """Add Stage 4 content"""
        doc.add_heading('STAGE 4: DEVELOPMENT OF THE BEST IDEA', level=1)
        
        stage_data = project_data.get('stage_4_data', {})
        
        doc.add_heading('4.1 Selected Solution', level=2)
        doc.add_paragraph(stage_data.get('development_details', 'Solution details pending.'))
    
    def _add_stage_5(self, doc: Document, project_data: Dict):
        """Add Stage 5 content"""
        doc.add_heading('STAGE 5: PRESENTATION OF RESULTS', level=1)
        
        stage_data = project_data.get('stage_5_data', {})
        
        doc.add_heading('5.1 Visual Presentation', level=2)
        generated_images = stage_data.get('generated_images', [])
        
        if generated_images:
            doc.add_paragraph(f"{len(generated_images)} visual aid(s) were created for this presentation.")
        else:
            doc.add_paragraph('Visual aids to be added.')
    
    def _add_stage_6(self, doc: Document, project_data: Dict):
        """Add Stage 6 content"""
        doc.add_heading('STAGE 6: EVALUATION AND RECOMMENDATIONS', level=1)
        
        stage_data = project_data.get('stage_6_data', {})
        
        doc.add_heading('6.1 Project Evaluation', level=2)
        evaluation = stage_data.get('evaluation', 'Evaluation pending.')
        doc.add_paragraph(evaluation)
        
        doc.add_heading('6.2 Recommendations', level=2)
        doc.add_paragraph('Based on the project findings, the following recommendations are proposed for future improvement and implementation.')
