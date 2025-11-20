"""
ZIMSEC Project Document DOCX Generator
Generates editable Word documents for ZIMSEC students
"""

import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class ProjectDocxGenerator:
    """Generates editable DOCX documents for ZIMSEC projects"""
    
    def __init__(self):
        # Ensure output directory exists
        self.output_dir = os.path.join('static', 'projects')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_project_docx(self, project_data: dict, user_id: str) -> str:
        """Generate complete ZIMSEC project document in DOCX format"""
        try:
            project_id = project_data.get('id', 'unknown')
            filename = f"ZIMSEC_Project_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(12)
            
            # Cover Page
            self._create_cover_page(doc, project_data)
            doc.add_page_break()
            
            # Table of Contents
            self._create_table_of_contents(doc)
            doc.add_page_break()
            
            # Main Content
            self._create_section_a(doc, project_data)
            doc.add_page_break()
            
            self._create_section_b(doc, project_data)
            doc.add_page_break()
            
            self._create_implementation_section(doc, project_data)
            doc.add_page_break()
            
            self._create_testing_section(doc, project_data)
            doc.add_page_break()
            
            self._create_documentation_section(doc, project_data)
            doc.add_page_break()
            
            self._create_evaluation_section(doc, project_data)
            
            # Save document
            doc.save(filepath)
            
            logger.info(f"✅ Generated DOCX document: {filename}")
            
            # Return public URL
            try:
                from utils.url_utils import convert_local_path_to_public_url
                return convert_local_path_to_public_url(filepath)
            except:
                return filepath
                
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise e
    
    def _add_heading(self, doc, text, level=1):
        """Add a formatted heading"""
        heading = doc.add_heading(text, level=level)
        if level == 1:
            heading.runs[0].font.color.rgb = RGBColor(156, 39, 176)  # Purple
        return heading
    
    def _create_cover_page(self, doc, project_data):
        """Create cover page"""
        data = project_data.get('project_data', {})
        
        # School name
        school = doc.add_paragraph(data.get('school', 'School Name').upper())
        school.alignment = WD_ALIGN_PARAGRAPH.CENTER
        school.runs[0].font.size = Pt(16)
        school.runs[0].font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Project title
        title = doc.add_paragraph(project_data.get('project_title', 'Project Title'))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(156, 39, 176)
        
        doc.add_paragraph()  # Spacing
        
        # Subject
        subject = doc.add_paragraph(f"Subject: {project_data.get('subject', 'Subject')}")
        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subject.runs[0].font.size = Pt(14)
        
        doc.add_paragraph()  # Spacing
        
        # Student details
        details = [
            f"Student Name: {data.get('student_name', '')} {data.get('student_surname', '')}",
            f"Form: {data.get('form', '')}",
            f"School: {data.get('school', '')}",
            f"Date: {datetime.now().strftime('%B %Y')}"
        ]
        
        for detail in details:
            p = doc.add_paragraph(detail)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _create_table_of_contents(self, doc):
        """Create table of contents"""
        self._add_heading(doc, "TABLE OF CONTENTS", level=1)
        
        toc_items = [
            "Section A: Selection, Investigation and Analysis",
            "Section B: Design",
            "Implementation",
            "Testing",
            "Documentation",
            "Evaluation"
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Bullet')
    
    def _create_section_a(self, doc, project_data):
        """Create Section A"""
        data = project_data.get('project_data', {})
        
        self._add_heading(doc, "SECTION A: SELECTION, INVESTIGATION AND ANALYSIS", level=1)
        
        self._add_heading(doc, "1. Problem Definition", level=2)
        doc.add_paragraph(data.get('problem_definition', '[Describe the problem your project addresses]'))
        
        self._add_heading(doc, "2. Investigation of Current System", level=2)
        doc.add_paragraph(data.get('investigation', '[Describe how the current system works]'))
        
        self._add_heading(doc, "3. Requirements", level=2)
        requirements = data.get('requirements', [])
        if isinstance(requirements, list):
            for req in requirements:
                doc.add_paragraph(req, style='List Bullet')
        else:
            doc.add_paragraph(str(requirements))
        
        self._add_heading(doc, "4. Aims and Objectives", level=2)
        objectives = data.get('objectives', [])
        if isinstance(objectives, list):
            for obj in objectives:
                doc.add_paragraph(obj, style='List Bullet')
        else:
            doc.add_paragraph(str(objectives))
    
    def _create_section_b(self, doc, project_data):
        """Create Section B"""
        data = project_data.get('project_data', {})
        
        self._add_heading(doc, "SECTION B: DESIGN", level=1)
        
        self._add_heading(doc, "1. Alternative Methods Considered", level=2)
        doc.add_paragraph(data.get('alternatives', '[Describe alternative approaches considered]'))
        
        self._add_heading(doc, "2. Input Design", level=2)
        doc.add_paragraph(data.get('input_design', '[Describe the inputs required]'))
        
        self._add_heading(doc, "3. Output Design", level=2)
        doc.add_paragraph(data.get('output_design', '[Describe the expected outputs]'))
        
        self._add_heading(doc, "4. Test Plan", level=2)
        doc.add_paragraph(data.get('test_plan', '[Describe your testing strategy]'))
    
    def _create_implementation_section(self, doc, project_data):
        """Create Implementation section"""
        data = project_data.get('project_data', {})
        
        self._add_heading(doc, "IMPLEMENTATION", level=1)
        doc.add_paragraph(data.get('implementation', '[Describe how you implemented the project]'))
    
    def _create_testing_section(self, doc, project_data):
        """Create Testing section"""
        data = project_data.get('project_data', {})
        
        self._add_heading(doc, "TESTING", level=1)
        doc.add_paragraph(data.get('testing', '[Document your testing process and results]'))
    
    def _create_documentation_section(self, doc, project_data):
        """Create Documentation section"""
        data = project_data.get('project_data', {})
        
        self._add_heading(doc, "DOCUMENTATION", level=1)
        
        self._add_heading(doc, "User Documentation", level=2)
        doc.add_paragraph(data.get('user_documentation', '[Instructions for users]'))
        
        self._add_heading(doc, "Technical Documentation", level=2)
        doc.add_paragraph(data.get('technical_documentation', '[Technical details for developers]'))
    
    def _create_evaluation_section(self, doc, project_data):
        """Create Evaluation section"""
        data = project_data.get('project_data', {})
        
        self._add_heading(doc, "EVALUATION", level=1)
        doc.add_paragraph(data.get('evaluation', '[Evaluate the success of your project]'))
        
        self._add_heading(doc, "Conclusion", level=2)
        doc.add_paragraph(data.get('conclusion', '[Summarize your project outcomes]'))
